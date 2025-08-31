#!/usr/bin/env python3
"""
Document Processing Pipeline for Semantic Search
==============================================

This module provides comprehensive document processing capabilities for the semantic search engine.
It builds upon the TextPreprocessor from the word embeddings module and extends it with advanced
chunking strategies and document format support.

Features:
- Multi-format document loading (text, PDF, markdown)
- Smart chunking strategies (fixed-size, sentence-boundary, semantic)
- Metadata extraction and management
- Integration with existing TextPreprocessor
- Batch processing capabilities

Author: AI Bootcamp Week 3-4
Date: 2025
"""

import os
import re
import json
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union, Any
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod

# Core libraries
import numpy as np

# Document processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF processing will be disabled.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX processing will be disabled.")

# Import the existing TextPreprocessor from word embeddings
import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..', '..', '02-word-embeddings'))

try:
    from word_embeddings_comprehensive import TextPreprocessor
    WORD_EMBEDDINGS_AVAILABLE = True
except ImportError:
    WORD_EMBEDDINGS_AVAILABLE = False
    print("Warning: Could not import TextPreprocessor from word embeddings module.")
    # Create a minimal fallback preprocessor
    class TextPreprocessor:
        def __init__(self, min_count=1):
            self.min_count = min_count
        
        def clean_text(self, text: str) -> str:
            text = text.lower()
            text = re.sub(r'[^a-z0-9\s]', ' ', text)
            text = ' '.join(text.split())
            return text


@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata."""
    chunk_id: str
    document_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for serialization."""
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'DocumentChunk':
        """Create chunk from dictionary."""
        return cls(**data)


@dataclass
class Document:
    """Represents a document with metadata."""
    document_id: str
    title: str
    content: str
    file_path: Optional[str]
    file_type: str
    metadata: Dict[str, Any]
    chunks: List[DocumentChunk]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary for serialization."""
        data = asdict(self)
        data['chunks'] = [chunk.to_dict() for chunk in self.chunks]
        return data
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'Document':
        """Create document from dictionary."""
        chunks_data = data.pop('chunks', [])
        doc = cls(**data)
        doc.chunks = [DocumentChunk.from_dict(chunk_data) for chunk_data in chunks_data]
        return doc


class DocumentLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def can_load(self, file_path: str) -> bool:
        """Check if this loader can handle the file type."""
        pass
    
    @abstractmethod
    def load(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load document content and extract metadata."""
        pass


class TextLoader(DocumentLoader):
    """Loader for plain text files."""
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is a text file."""
        return file_path.lower().endswith(('.txt', '.md', '.markdown', '.rst'))
    
    def load(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                'file_size': os.path.getsize(file_path),
                'encoding': 'utf-8'
            }
            
            return content, metadata
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    metadata = {
                        'file_size': os.path.getsize(file_path),
                        'encoding': encoding
                    }
                    
                    return content, metadata
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode file {file_path} with any supported encoding")


class PDFLoader(DocumentLoader):
    """Loader for PDF files."""
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is a PDF and PyPDF2 is available."""
        return PDF_AVAILABLE and file_path.lower().endswith('.pdf')
    
    def load(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load PDF file."""
        if not PDF_AVAILABLE:
            raise RuntimeError("PyPDF2 not installed. Cannot load PDF files.")
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                content_parts = []
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(page_text)
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                
                content = '\n\n'.join(content_parts)
                
                metadata = {
                    'num_pages': len(reader.pages),
                    'file_size': os.path.getsize(file_path)
                }
                
                # Extract additional metadata if available
                if reader.metadata:
                    metadata.update({
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                        'creator': reader.metadata.get('/Creator', '')
                    })
                
                return content, metadata
                
        except Exception as e:
            raise ValueError(f"Error loading PDF file {file_path}: {e}")


class ChunkingStrategy(ABC):
    """Abstract base class for text chunking strategies."""
    
    @abstractmethod
    def chunk_text(self, text: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk text into segments."""
        pass


class FixedSizeChunker(ChunkingStrategy):
    """Chunks text into fixed-size segments with optional overlap."""
    
    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        """
        Initialize fixed-size chunker.
        
        Args:
            chunk_size: Target size of each chunk in characters
            overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk text into fixed-size segments."""
        chunks = []
        text_length = len(text)
        
        if text_length == 0:
            return chunks
        
        start = 0
        chunk_index = 0
        
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            
            # Extract chunk content
            chunk_content = text[start:end].strip()
            
            if chunk_content:  # Only create chunk if it has content
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                
                chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=chunk_content,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    metadata={
                        **metadata,
                        'chunk_type': 'fixed_size',
                        'chunk_size_actual': len(chunk_content)
                    }
                )
                
                chunks.append(chunk)
                chunk_index += 1
            
            # Move to next chunk with overlap
            start = end - self.overlap
            
            # Prevent infinite loop
            if start >= end:
                start = end
        
        return chunks


class SentenceBoundaryChunker(ChunkingStrategy):
    """Chunks text at sentence boundaries to preserve semantic coherence."""
    
    def __init__(self, target_size: int = 512, max_size: int = 768):
        """
        Initialize sentence boundary chunker.
        
        Args:
            target_size: Target size of each chunk in characters
            max_size: Maximum size before forcing a split
        """
        self.target_size = target_size
        self.max_size = max_size
        
        # Sentence boundary patterns
        self.sentence_endings = re.compile(r'[.!?]+\s+')
    
    def chunk_text(self, text: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk text at sentence boundaries."""
        chunks = []
        
        if not text.strip():
            return chunks
        
        # Split into sentences
        sentences = self._split_sentences(text)
        
        current_chunk = []
        current_length = 0
        chunk_index = 0
        start_char = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed target size
            if current_chunk and (current_length + sentence_length > self.target_size):
                # Create chunk from current sentences
                chunk_content = ' '.join(current_chunk).strip()
                
                if chunk_content:
                    chunk_id = f"{document_id}_chunk_{chunk_index}"
                    end_char = start_char + len(chunk_content)
                    
                    chunk = DocumentChunk(
                        chunk_id=chunk_id,
                        document_id=document_id,
                        content=chunk_content,
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=end_char,
                        metadata={
                            **metadata,
                            'chunk_type': 'sentence_boundary',
                            'num_sentences': len(current_chunk),
                            'chunk_size_actual': len(chunk_content)
                        }
                    )
                    
                    chunks.append(chunk)
                    chunk_index += 1
                    start_char = end_char + 1  # Account for space
                
                # Start new chunk
                current_chunk = []
                current_length = 0
            
            # Handle very long sentences
            if sentence_length > self.max_size:
                # If we have accumulated sentences, create a chunk first
                if current_chunk:
                    chunk_content = ' '.join(current_chunk).strip()
                    if chunk_content:
                        chunk_id = f"{document_id}_chunk_{chunk_index}"
                        end_char = start_char + len(chunk_content)
                        
                        chunk = DocumentChunk(
                            chunk_id=chunk_id,
                            document_id=document_id,
                            content=chunk_content,
                            chunk_index=chunk_index,
                            start_char=start_char,
                            end_char=end_char,
                            metadata={
                                **metadata,
                                'chunk_type': 'sentence_boundary',
                                'num_sentences': len(current_chunk),
                                'chunk_size_actual': len(chunk_content)
                            }
                        )
                        
                        chunks.append(chunk)
                        chunk_index += 1
                        start_char = end_char + 1
                    
                    current_chunk = []
                    current_length = 0
                
                # Split long sentence using fixed-size chunker
                long_sentence_chunks = FixedSizeChunker(self.max_size, 50).chunk_text(
                    sentence, document_id, metadata
                )
                
                for sub_chunk in long_sentence_chunks:
                    sub_chunk.chunk_id = f"{document_id}_chunk_{chunk_index}"
                    sub_chunk.chunk_index = chunk_index
                    sub_chunk.start_char = start_char
                    sub_chunk.end_char = start_char + len(sub_chunk.content)
                    sub_chunk.metadata['chunk_type'] = 'sentence_boundary_long'
                    
                    chunks.append(sub_chunk)
                    chunk_index += 1
                    start_char = sub_chunk.end_char + 1
            else:
                # Add sentence to current chunk
                current_chunk.append(sentence)
                current_length += sentence_length
        
        # Create final chunk if there are remaining sentences
        if current_chunk:
            chunk_content = ' '.join(current_chunk).strip()
            
            if chunk_content:
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                end_char = start_char + len(chunk_content)
                
                chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=chunk_content,
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=end_char,
                    metadata={
                        **metadata,
                        'chunk_type': 'sentence_boundary',
                        'num_sentences': len(current_chunk),
                        'chunk_size_actual': len(chunk_content)
                    }
                )
                
                chunks.append(chunk)
        
        return chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting - can be enhanced with more sophisticated methods
        sentences = []
        
        # Split on sentence endings
        parts = self.sentence_endings.split(text)
        
        current_sentence = ""
        for i, part in enumerate(parts):
            current_sentence += part
            
            # If this is not the last part, we had a sentence ending
            if i < len(parts) - 1:
                # Find the actual ending that was matched
                remaining_text = text[len(''.join(parts[:i+1])):]
                match = self.sentence_endings.match(remaining_text)
                if match:
                    current_sentence += match.group().strip()
                
                # Add sentence if it's not empty
                if current_sentence.strip():
                    sentences.append(current_sentence.strip())
                current_sentence = ""
        
        # Add any remaining text
        if current_sentence.strip():
            sentences.append(current_sentence.strip())
        
        return [s for s in sentences if s.strip()]


class DocumentProcessor:
    """Main document processing pipeline that orchestrates loading, chunking, and preprocessing."""
    
    def __init__(self, preprocessor: Optional[TextPreprocessor] = None, 
                 chunking_strategy: Optional[ChunkingStrategy] = None):
        """
        Initialize document processor.
        
        Args:
            preprocessor: Text preprocessor instance (uses existing one from word embeddings)
            chunking_strategy: Strategy for chunking documents
        """
        self.preprocessor = preprocessor or TextPreprocessor(min_count=1)
        self.chunking_strategy = chunking_strategy or SentenceBoundaryChunker()
        
        # Initialize document loaders
        self.loaders = [
            TextLoader(),
            PDFLoader()
        ]
        
        # Add DOCX loader if available
        if DOCX_AVAILABLE:
            self.loaders.append(DOCXLoader())
    
    def process_document(self, file_path: str, document_id: Optional[str] = None) -> Document:
        """
        Process a single document file.
        
        Args:
            file_path: Path to the document file
            document_id: Optional custom document ID
        
        Returns:
            Processed Document object with chunks
        """
        file_path = str(Path(file_path).resolve())
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Find appropriate loader
        loader = None
        for l in self.loaders:
            if l.can_load(file_path):
                loader = l
                break
        
        if loader is None:
            raise ValueError(f"No loader available for file type: {file_path}")
        
        # Load document content
        try:
            content, file_metadata = loader.load(file_path)
        except Exception as e:
            raise ValueError(f"Error loading document {file_path}: {e}")
        
        # Generate document ID if not provided
        if document_id is None:
            document_id = Path(file_path).stem
        
        # Extract file information
        file_info = Path(file_path)
        file_type = file_info.suffix.lower().lstrip('.')
        title = file_info.stem
        
        # Combine metadata
        metadata = {
            'created_at': os.path.getctime(file_path),
            'modified_at': os.path.getmtime(file_path),
            'file_extension': file_type,
            **file_metadata
        }
        
        # Create document object
        document = Document(
            document_id=document_id,
            title=title,
            content=content,
            file_path=file_path,
            file_type=file_type,
            metadata=metadata,
            chunks=[]
        )
        
        # Chunk the document
        chunks = self.chunking_strategy.chunk_text(content, document_id, metadata)
        document.chunks = chunks
        
        return document
    
    def process_text(self, text: str, document_id: str, title: Optional[str] = None,
                    metadata: Optional[Dict[str, Any]] = None) -> Document:
        """
        Process raw text content.
        
        Args:
            text: Text content to process
            document_id: Unique identifier for the document
            title: Optional title for the document
            metadata: Optional metadata dictionary
        
        Returns:
            Processed Document object with chunks
        """
        if metadata is None:
            metadata = {}
        
        # Create document object
        document = Document(
            document_id=document_id,
            title=title or document_id,
            content=text,
            file_path=None,
            file_type='text',
            metadata={
                'content_length': len(text),
                'created_from': 'text_input',
                **metadata
            },
            chunks=[]
        )
        
        # Chunk the document
        chunks = self.chunking_strategy.chunk_text(text, document_id, document.metadata)
        document.chunks = chunks
        
        return document
    
    def process_directory(self, directory_path: str, 
                         file_patterns: Optional[List[str]] = None,
                         recursive: bool = False) -> List[Document]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            file_patterns: Optional list of file patterns to match
            recursive: Whether to process subdirectories recursively
        
        Returns:
            List of processed Document objects
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if not directory_path.is_dir():
            raise ValueError(f"Path is not a directory: {directory_path}")
        
        # Collect files to process
        files_to_process = []
        
        if recursive:
            for pattern in (file_patterns or ['**/*']):
                files_to_process.extend(directory_path.glob(pattern))
        else:
            for pattern in (file_patterns or ['*']):
                files_to_process.extend(directory_path.glob(pattern))
        
        # Filter for supported file types and actual files
        supported_files = []
        for file_path in files_to_process:
            if file_path.is_file():
                # Check if any loader can handle this file
                for loader in self.loaders:
                    if loader.can_load(str(file_path)):
                        supported_files.append(file_path)
                        break
        
        # Process each file
        documents = []
        for file_path in supported_files:
            try:
                document = self.process_document(str(file_path))
                documents.append(document)
                print(f"Processed: {file_path.name} -> {len(document.chunks)} chunks")
            except Exception as e:
                print(f"Error processing {file_path.name}: {e}")
        
        print(f"Successfully processed {len(documents)} documents from {directory_path}")
        return documents
    
    def get_processing_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        Get statistics about processed documents.
        
        Args:
            documents: List of processed documents
        
        Returns:
            Dictionary containing processing statistics
        """
        if not documents:
            return {'total_documents': 0}
        
        total_chunks = sum(len(doc.chunks) for doc in documents)
        total_chars = sum(len(doc.content) for doc in documents)
        chunk_sizes = [len(chunk.content) for doc in documents for chunk in doc.chunks]
        
        file_types = {}
        for doc in documents:
            file_types[doc.file_type] = file_types.get(doc.file_type, 0) + 1
        
        stats = {
            'total_documents': len(documents),
            'total_chunks': total_chunks,
            'total_characters': total_chars,
            'average_chunks_per_doc': total_chunks / len(documents),
            'average_doc_length': total_chars / len(documents),
            'file_types': file_types,
            'chunk_size_stats': {
                'min': min(chunk_sizes) if chunk_sizes else 0,
                'max': max(chunk_sizes) if chunk_sizes else 0,
                'mean': np.mean(chunk_sizes) if chunk_sizes else 0,
                'std': np.std(chunk_sizes) if chunk_sizes else 0
            }
        }
        
        return stats


# Optional DOCX loader (only if python-docx is available)
if DOCX_AVAILABLE:
    class DOCXLoader(DocumentLoader):
        """Loader for Microsoft Word documents."""
        
        def can_load(self, file_path: str) -> bool:
            """Check if file is a DOCX file."""
            return file_path.lower().endswith('.docx')
        
        def load(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
            """Load DOCX file."""
            try:
                doc = docx.Document(file_path)
                
                # Extract text from paragraphs
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                
                content = '\n'.join(content_parts)
                
                # Extract metadata
                metadata = {
                    'file_size': os.path.getsize(file_path),
                    'num_paragraphs': len([p for p in doc.paragraphs if p.text.strip()])
                }
                
                # Core properties
                core_props = doc.core_properties
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.created:
                    metadata['created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['modified'] = core_props.modified.isoformat()
                
                return content, metadata
                
            except Exception as e:
                raise ValueError(f"Error loading DOCX file {file_path}: {e}")


def create_sample_processor(chunking_method: str = "sentence") -> DocumentProcessor:
    """
    Create a document processor with a specific chunking strategy.
    
    Args:
        chunking_method: Either "fixed" or "sentence"
    
    Returns:
        Configured DocumentProcessor instance
    """
    preprocessor = TextPreprocessor(min_count=1)
    
    if chunking_method == "fixed":
        chunker = FixedSizeChunker(chunk_size=512, overlap=50)
    elif chunking_method == "sentence":
        chunker = SentenceBoundaryChunker(target_size=512, max_size=768)
    else:
        raise ValueError(f"Unknown chunking method: {chunking_method}")
    
    return DocumentProcessor(preprocessor=preprocessor, chunking_strategy=chunker)


if __name__ == "__main__":
    # Simple test
    print("Document Processor Test")
    print("=" * 50)
    
    # Test with sample text
    sample_text = """
    Artificial intelligence (AI) is intelligence demonstrated by machines, in contrast to 
    the natural intelligence displayed by humans and animals. Leading AI textbooks define 
    the field as the study of "intelligent agents": any device that perceives its environment 
    and takes actions that maximize its chance of successfully achieving its goals.
    
    Machine learning is a subset of AI that provides systems the ability to automatically 
    learn and improve from experience without being explicitly programmed. Machine learning 
    focuses on the development of computer programs that can access data and use it to learn 
    for themselves.
    
    Deep learning is part of a broader family of machine learning methods based on artificial 
    neural networks with representation learning. Learning can be supervised, semi-supervised 
    or unsupervised.
    """
    
    # Create processor
    processor = create_sample_processor("sentence")
    
    # Process text
    document = processor.process_text(
        text=sample_text,
        document_id="ai_overview",
        title="AI and Machine Learning Overview",
        metadata={"topic": "artificial_intelligence", "source": "example"}
    )
    
    print(f"Document: {document.title}")
    print(f"Content length: {len(document.content)} characters")
    print(f"Number of chunks: {len(document.chunks)}")
    print()
    
    # Display chunks
    for i, chunk in enumerate(document.chunks):
        print(f"Chunk {i + 1} (ID: {chunk.chunk_id}):")
        print(f"  Length: {len(chunk.content)} characters")
        print(f"  Content: {chunk.content[:100]}...")
        if len(chunk.content) > 100:
            print(f"    [...{chunk.content[-50:]}]")
        print()